# -*- coding: utf-8 -*-
"""
生成《南宁师范大学本科毕业论文（设计）开题报告》docx
按学校写作指南格式规范：正文小四宋体/固定行距25磅/首行缩进2字符
一级标题四号黑体，二级标题小四黑体，三级标题小四仿宋加粗
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_font(run, name_east, size, bold=False, name_west='Times New Roman'):
    """设置 run 的中西文字体、字号、加粗"""
    run.font.name = name_west
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name_east)


def add_para(doc, text, font='宋体', size=12, bold=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent_chars=2,
             line_exact=25, space_before=0, space_after=0):
    """添加一个正文段落，行距为固定值(磅)，支持首行缩进(字符数)"""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_exact)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent_chars:
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLineChars'), str(first_indent_chars * 100))
        ind.set(qn('w:firstLine'), str(int(first_indent_chars * size * 20)))
    run = p.add_run(text)
    set_run_font(run, font, size, bold)
    return p


def add_heading1(doc, text):
    """一级标题：四号黑体，段前1行段后0.5行，左对齐缩进2字符"""
    return add_para(doc, text, font='黑体', size=14, bold=True,
                    align=WD_ALIGN_PARAGRAPH.LEFT, first_indent_chars=2,
                    line_exact=25, space_before=14, space_after=7)


def add_heading2(doc, text):
    """二级标题：小四黑体"""
    return add_para(doc, text, font='黑体', size=12, bold=True,
                    align=WD_ALIGN_PARAGRAPH.LEFT, first_indent_chars=2,
                    line_exact=25, space_before=0, space_after=0)


def add_heading3(doc, text):
    """三级标题：小四仿宋加粗"""
    return add_para(doc, text, font='仿宋', size=12, bold=True,
                    align=WD_ALIGN_PARAGRAPH.LEFT, first_indent_chars=2,
                    line_exact=25, space_before=0, space_after=0)


def set_cell_text(cell, text, font='宋体', size=12, bold=False,
                  align=WD_ALIGN_PARAGRAPH.CENTER):
    """设置表格单元格文字"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(22)
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, font, size, bold)


def set_cell_shading(cell, color_hex):
    """给单元格填充底色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


doc = Document()

# 页面设置 A4
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.54)
sec.bottom_margin = Cm(2.54)
sec.left_margin = Cm(3.0)
sec.right_margin = Cm(3.0)

# 默认样式兜底
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============ 标题 ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = p.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
pf.line_spacing = Pt(30)
pf.space_before = Pt(0)
pf.space_after = Pt(10)
run = p.add_run('南宁师范大学本科毕业论文（设计）开题报告')
set_run_font(run, '黑体', 16, True)

# ============ 基本信息表 ============
t = doc.add_table(rows=5, cols=4)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
# 第一行：论文题目（合并4列，跨整行）
a = t.cell(0, 0)
b = t.cell(0, 1)
c = t.cell(0, 2)
d = t.cell(0, 3)
a.merge(b).merge(c).merge(d)
set_cell_text(a, '论文题目：基于RAG的智能学习系统的设计与实现', '黑体', 12, True, WD_ALIGN_PARAGRAPH.LEFT)
set_cell_shading(a, 'F2F2F2')
# 其余行
rest = [
    ('学生姓名', '黄朕', '学    号', '2515100120'),
    ('专业班级', '软件工程（计23专升本1班）', '学    院', '计算机与信息工程学院'),
    ('指导教师', '　　　　　　　　', '职    称', '　　　　　　　　'),
    ('开题日期', '2026年　　月　　日', '计划完成时间', '2026年　　月　　日'),
]
for i, (a, b, c_, d) in enumerate(rest, start=1):
    set_cell_text(t.cell(i, 0), a, '黑体', 12, True)
    set_cell_shading(t.cell(i, 0), 'F2F2F2')
    set_cell_text(t.cell(i, 1), b, '宋体', 12, False, WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(t.cell(i, 2), c_, '黑体', 12, True)
    set_cell_shading(t.cell(i, 2), 'F2F2F2')
    set_cell_text(t.cell(i, 3), d, '宋体', 12, False, WD_ALIGN_PARAGRAPH.LEFT)
for row in t.rows:
    row.cells[0].width = Cm(2.8)
    row.cells[1].width = Cm(4.6)
    row.cells[2].width = Cm(2.8)
    row.cells[3].width = Cm(4.6)

# ============ 一、选题的依据与意义 ============
add_heading1(doc, '一、选题的依据与意义')

add_heading2(doc, '（一）选题背景')
add_para(doc, '随着人工智能技术的迅猛发展，大规模语言模型（Large Language Model, LLM）在自然语言理解与生成方面展现出强大的能力，为教育信息化带来了新的契机。然而，传统学习平台普遍存在以下问题：')
add_heading3(doc, '1. AI幻觉问题')
add_para(doc, '通用大模型在回答课程相关知识时，容易生成与教材不符甚至错误的内容，缺乏对具体课程材料的约束，回答的可信度难以保证。')
add_heading3(doc, '2. 交互方式单一')
add_para(doc, '大多数学习App仅提供课程浏览、视频播放与固定题库练习，缺乏基于自然语言的智能问答与个性化学习反馈，难以满足学生个性化学习需求。')
add_heading3(doc, '3. 管理手段落后')
add_para(doc, '职业院校的排课、调课、考试发布仍依赖人工或老旧系统，效率低、易冲突，缺乏信息化、自动化的排课支持。')
add_heading3(doc, '4. 评估不够个性化')
add_para(doc, '测评往往只给出单一分数，缺少对学习者多维能力（逻辑思维、专注耐力、信息检索等）的刻画与自适应引导，无法有效支持因材施教。')
add_para(doc, '在此背景下，本课题设计并实现一个基于检索增强生成（Retrieval-Augmented Generation, RAG）的智能学习系统，将大模型能力与校本课程知识库有机结合，面向职业院校师生提供多端一体的智能学习与管理服务。')

add_heading2(doc, '（二）选题意义')
add_heading3(doc, '1. 理论意义')
add_para(doc, '本课题将RAG技术引入垂直领域教育场景，探索如何通过课程文档向量化检索缓解大模型幻觉、实现“可溯源”的智能问答，具有一定的应用研究价值；同时将智能排课、个性化测评等算法融入系统，丰富了教育信息化的技术实践。')
add_heading3(doc, '2. 实践意义')
add_para(doc, '系统覆盖Android学生端、Web管理端与微信小程序三端，实现课程管理、智能排课、班级聊天、@AI问答、多维度刷题测评、考试作业发布等完整业务闭环，可直接服务于职业院校的教学管理，具备较强的实用性与推广价值。')

# ============ 二、国内外研究现状 ============
add_heading1(doc, '二、国内外研究现状')

add_heading2(doc, '（一）国外研究现状')
add_para(doc, '国际上，检索增强生成（RAG）由Meta等机构于2020年提出，通过“先检索、后生成”的方式将外部知识库注入大模型，显著降低幻觉并提升回答可解释性，已成为大模型落地的核心技术范式。国外教育平台如Khan Academy的Khanmigo已探索将大模型用于个性化辅导；同时，基于向量数据库（如Milvus、Pinecone）的语义检索技术日趋成熟，为教育知识库问答提供了可靠的工程基础。')

add_heading2(doc, '（二）国内研究现状')
add_para(doc, '国内方面，超星学习通、雨课堂、智慧职教等平台已普遍应用，但多集中在资源分发与课堂管理，AI交互与个性化测评仍处于探索阶段。近年来，国内大模型（如通义千问、DeepSeek等）与开源向量化模型（如BGE系列）的快速发展，为高校自主构建“课程知识库+大模型”的智能问答系统提供了低成本、可本地化的技术路径，相关研究在医疗、政务、教育等行业加速落地。')

add_heading2(doc, '（三）现状总结')
add_para(doc, '现有研究成果为本课题提供了理论支撑，但面向职业院校的、融合RAG智能问答与智能排课、且覆盖多端的一体化学习系统仍较少，本课题正是在这一方向上开展的设计与实现工作。')

# ============ 三、主要研究内容 ============
add_heading1(doc, '三、主要研究内容')
add_para(doc, '本课题设计并实现一套基于RAG的智能学习系统，主要研究内容包括：')
add_heading3(doc, '1. RAG智能问答模块')
add_para(doc, '课程文档（Word/PDF）解析、文本分块、向量化（BGE-M3）、余弦相似度检索（Top-5）与LLM生成，并在检索失败时降级为关键词匹配，实现“可溯源”的课程问答。')
add_heading3(doc, '2. 智能排课与调课模块')
add_para(doc, '基于贪心加回溯算法的自动排课，结合BitSet位图实现O(1)时间复杂度的课时冲突检测，支持按周独立排课、跨周调课与课程锁定等场景。')
add_heading3(doc, '3. 多维度答题与测评模块')
add_para(doc, 'AI智能出题，作答后生成六维能力评分（逻辑思维、判断决策、专注耐力、专业学习力、信息检索、自律执行），并实现自适应难度与错题分析缓存。')
add_heading3(doc, '4. 实时聊天与通知模块')
add_para(doc, '支持文字、图片、文件、链接消息，@学生定向发送，@AI智能应答，以及基于WebSocket的实时推送与多级消息去重。')
add_heading3(doc, '5. 多端一体化系统')
add_para(doc, 'Android学生端、Web管理端（React）、微信小程序三端共用后端服务，实现统一认证与数据管理。')

# ============ 四、研究方法与技术路线 ============
add_heading1(doc, '四、研究方法与技术路线')

add_heading2(doc, '（一）研究方法')
add_heading3(doc, '1. 文献研究法')
add_para(doc, '查阅RAG、大模型、智能排课、教育信息化等领域文献，确定技术方案与可行性。')
add_heading3(doc, '2. 对比分析法')
add_para(doc, '对比传统关键词检索与向量检索、对比直接调用大模型与RAG问答的效果差异，论证技术选型。')
add_heading3(doc, '3. 系统开发法')
add_para(doc, '采用软件工程流程，完成需求分析、概要设计、详细设计、编码实现与测试。')
add_heading3(doc, '4. 测试验证法')
add_para(doc, '面向管理员、教师、学生三类角色设计功能测试用例与边界测试，验证系统正确性与稳定性。')

add_heading2(doc, '（二）技术路线')
add_para(doc, '系统采用前后端分离架构，总体技术栈如下：')
tech_rows = [
    ('系统层次', '技术选型'),
    ('后端', 'Java 17、Spring Boot 3.2、MyBatis-Plus、MySQL、Spring Security、WebSocket'),
    ('AI层', 'RAG（BGE-M3向量化 + 余弦相似度检索 + LLM生成）'),
    ('学生端', 'Android（Retrofit、OkHttp、ViewPager2）'),
    ('管理端', 'React、TypeScript、Vite、Recharts'),
    ('小程序端', 'uni-app'),
    ('安全', 'JWT无状态认证、BCrypt密码加密、RBAC权限控制'),
]
tt = doc.add_table(rows=len(tech_rows), cols=2)
tt.style = 'Table Grid'
tt.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (a, b) in enumerate(tech_rows):
    set_cell_text(tt.cell(i, 0), a, '黑体', 11, True)
    set_cell_shading(tt.cell(i, 0), 'F2F2F2')
    set_cell_text(tt.cell(i, 1), b, '宋体', 11, False, WD_ALIGN_PARAGRAPH.LEFT)
for row in tt.rows:
    row.cells[0].width = Cm(2.8)
    row.cells[1].width = Cm(12.0)
add_para(doc, '', first_indent_chars=0, line_exact=12)
add_para(doc, '技术路线流程为：需求分析 → 数据库设计 → 后端接口开发 → AI检索与测评逻辑开发 → 三端前端开发 → 联调测试 → 部署演示。')

# ============ 五、进度安排 ============
add_heading1(doc, '五、进度安排')
plan_rows = [
    ('阶段', '时间', '主要任务'),
    ('1', '第1-2周', '文献调研、选题论证、撰写开题报告'),
    ('2', '第3-4周', '需求分析、数据库设计、总体架构设计'),
    ('3', '第5-8周', '后端核心模块开发（认证、排课、RAG、测评、聊天）'),
    ('4', '第9-11周', 'Android学生端、Web管理端、小程序开发与对接'),
    ('5', '第12-13周', '系统集成、功能测试与边界测试、缺陷修复'),
    ('6', '第14-15周', '撰写毕业论文、整理图表'),
    ('7', '第16周', '论文修改、查重、答辩准备'),
]
tt = doc.add_table(rows=len(plan_rows), cols=3)
tt.style = 'Table Grid'
tt.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (a, b, c_) in enumerate(plan_rows):
    set_cell_text(tt.cell(i, 0), a, '黑体', 11, True)
    set_cell_shading(tt.cell(i, 0), 'F2F2F2')
    set_cell_text(tt.cell(i, 1), b, '宋体', 11, False)
    set_cell_text(tt.cell(i, 2), c_, '宋体', 11, False, WD_ALIGN_PARAGRAPH.LEFT)
for row in tt.rows:
    row.cells[0].width = Cm(2.0)
    row.cells[1].width = Cm(3.4)
    row.cells[2].width = Cm(9.4)

# ============ 六、预期成果 ============
add_heading1(doc, '六、预期成果')
add_heading3(doc, '1. 系统成果')
add_para(doc, '一套可运行的多端智能学习系统（Android学生端 + Web管理端 + 微信小程序）。')
add_heading3(doc, '2. 功能成果')
add_para(doc, '实现RAG智能问答、智能排课、六维测评、实时聊天等核心功能。')
add_heading3(doc, '3. 质量成果')
add_para(doc, '完成功能测试与边界测试，系统稳定运行。')
add_heading3(doc, '4. 论文成果')
add_para(doc, '撰写毕业论文一篇，完成毕业答辩。')

# ============ 参考文献 ============
add_heading1(doc, '参考文献')
refs = [
    '[1] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//Advances in Neural Information Processing Systems. 2020.',
    '[2] Gao Y, Xiong Y, Gao X, et al. Retrieval-Augmented Generation for Large Language Models: A Survey[J]. arXiv preprint arXiv:2312.10997, 2023.',
    '[3] 李刚. 疯狂Spring Boot讲义[M]. 北京: 电子工业出版社, 2020.',
    '[4] 王福强. 分布式系统架构：技术栈详解与快速进阶[M]. 北京: 机械工业出版社, 2019.',
    '[5] 郑阿奇. MySQL数据库原理与应用[M]. 北京: 电子工业出版社, 2018.',
    '[6] 曹雪虹, 张宗橙. 信息论与编码[M]. 北京: 清华大学出版社, 2016.',
    '[7] Amatriain X. Transformer Models: An Introduction and Catalog[J]. arXiv preprint arXiv:2302.07730, 2023.',
]
for ref in refs:
    add_para(doc, ref, font='宋体', size=10.5, first_indent_chars=0,
             align=WD_ALIGN_PARAGRAPH.LEFT, line_exact=20)

# ============ 指导教师意见 ============
add_heading1(doc, '指导教师意见')
add_para(doc, '本课题选题紧密结合教育信息化与人工智能发展趋势，技术路线清晰，研究内容具体，工作量饱满，具备较强的理论意义与实践价值，同意开题。')
add_para(doc, '', first_indent_chars=0, line_exact=25)
add_para(doc, '指导教师签名：　　　　　　　　　　　日期：　　　　年　　月　　日', first_indent_chars=0)

# ============ 学院审核意见 ============
add_heading1(doc, '学院审核意见')
add_para(doc, '', first_indent_chars=0, line_exact=25)
add_para(doc, '负责人签字：　　　　　　　　　　　　日期：　　　　年　　月　　日', first_indent_chars=0)

out = r'C:\Users\jay\Desktop\myBishe\aiStudy\南宁师范大学本科毕业论文开题报告_黄朕.docx'
doc.save(out)
print('已生成：', out)
